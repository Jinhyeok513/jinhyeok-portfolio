from __future__ import annotations

import shutil
import subprocess
import tempfile
import os
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
CAREER = ROOT / "docs" / "career"
PUBLIC_RESUME = ROOT / "public" / "resume"
SOFFICE = Path("/Users/jin/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/override/soffice")


CONTACT = {
    "name": "Jinhyeok Kim",
    "headline": "Artificial Intelligence Graduate",
    "phone": "0468 544 027",
    "email": "winsuwon127@gmail.com",
    "location": "Sydney, NSW, Australia",
    "portfolio": "jinhyeok-portfolio-amber.vercel.app",
    "github": "github.com/Jinhyeok513",
}

COURSEWORK = [
    "Introduction to Data Analytics",
    "Advanced Data Analytics Algorithms, Machine Learning",
    "Design, Data, and Decisions",
    "Project Management and the Professional",
    "AI/Analytics Capstone Project",
    "Communication for IT Professionals",
    "Data Structures and Algorithms",
    "Image Processing and Pattern Recognition",
    "Deep Learning and Convolutional Neural Network",
    "Introduction to Computational Intelligence",
    "Introduction to Artificial Intelligence",
    "The Ethics of Data and AI",
]

EDUCATION = [
    {
        "degree": "Bachelor of Artificial Intelligence",
        "school": "University of Technology Sydney (UTS), Sydney, Australia",
        "period": "Completed 2026",
        "details": [
            "Graduated with Distinction; WAM 77.94 and GPA 5.88/7.00.",
            "Relevant study included data analytics, machine learning, computer vision, generative AI, software development and business information systems.",
        ],
    },
    {
        "degree": "Diploma of Information Technology",
        "school": "UTS College, Sydney, Australia",
        "period": "2023",
        "details": [
            "Built foundations in Python and Java programming, web systems, system analysis, technical documentation and business information systems.",
        ],
    },
]

LANGUAGES = [
    "Korean - Native",
    "English - Professional communication skills",
    "Chinese - HSK Level 4",
]

WORK_RIGHTS = (
    "Current Australian work-rights details and full-time availability available on request."
)


@dataclass
class DocSpec:
    title: str
    subtitle: str
    folder: str
    filename: str
    kind: str
    summary: str
    skills: list[str]
    project_order: list[str]
    project_bullets: dict[str, list[str]]
    cv_objective: str = ""


PROJECTS = {
    "BioGeoDA": {
        "title": "BioGeoDA - Australian Plant Data Integration and AI-Assisted Trait Extraction",
        "period": "2025-2026",
        "context": "UTS AI/Analytics capstone team project and public portfolio reconstruction",
        "tools": "Python | Pandas | NLP | TF-IDF | Logistic Regression | BERT QA | Streamlit | GitHub",
        "bullets": [
            "Cleaned and standardised multi-source plant datasets by reviewing missing values, duplicate records, inconsistent species names and mixed trait formats.",
            "Developed keyword-tagging and fuzzy-matching workflows to locate species in OCR-processed botanical text while retaining source, page, paragraph and context metadata.",
            "Worked across TF-IDF classification and BERT-based question-answering pipelines for keyword-dependent and context-dependent trait extraction.",
            "Contributed to validating 844 new trait records from more than 2,400 AI-generated candidates across a 1,674-species project checklist.",
        ],
    },
    "AI Gym Trainer": {
        "title": "AI Gym Trainer - Exercise Video Analysis Application",
        "period": "2025-2026",
        "context": "Portfolio AI application",
        "tools": "Python | Next.js | FastAPI | MobileNetV2 | MediaPipe Pose | OpenCV | CSV/JSON export",
        "bullets": [
            "Built an end-to-end video analysis workflow covering upload, preprocessing, exercise classification, pose extraction, annotated video output and session reporting.",
            "Integrated MobileNetV2 classification with MediaPipe Pose landmarks, OpenCV rendering, joint-angle calculations and rule-based feedback.",
            "Documented model evidence honestly: MobileNetV2 reached 72.0% video accuracy and 78.18% video macro F1 on held-out comparison data, while a 10-video field test showed domain-shift limitations.",
        ],
    },
    "Tennis": {
        "title": "Tennis Ball Tracking and Trajectory Visualisation",
        "period": "2025",
        "context": "Computer-vision portfolio project",
        "tools": "Python | OpenCV | NumPy | Kalman filtering | FFmpeg | IoU/CLE evaluation",
        "bullets": [
            "Implemented a Kalman-filter tracking pipeline using HSV colour cues, whiteness cues, optical flow, blobness scoring, court masking and net-region suppression.",
            "Packaged a 207-frame tracking output with representative frame captures, trajectory visualisation and frame-level diagnostics.",
            "Reported sparse-label evaluation limits transparently, including mean IoU 0.1827, Success@IoU>=0.5 of 2.4% and mean CLE 131.18 px.",
        ],
    },
    "Project Management": {
        "title": "Event Management System Project Plan - Team Lead",
        "period": "2025",
        "context": "Academic Project Management Case Study",
        "tools": "Project Management | Excel | WBS | AON | Gantt scheduling | Risk and budget planning",
        "bullets": [
            "Led a seven-member academic team in developing an end-to-end project plan for a proposed AUD 1.5 million Event Management System supporting a major Sydney festival scenario.",
            "Directed task allocation, reviewed team deliverables and integrated sections covering scope, stakeholders, risk, quality, budget, resources and change control.",
            "Coordinated an Excel-based schedule and WBS with more than 100 tasks, dependencies, milestones, durations, float and critical-path analysis across an 18-month plan.",
        ],
    },
}


MASTER_SKILLS = [
    "Python", "Pandas", "NumPy", "SQL", "Data cleaning", "Data validation",
    "Machine learning", "NLP", "Computer vision", "Microsoft Excel - Intermediate",
    "PivotTables", "XLOOKUP/VLOOKUP", "SUMIFS/COUNTIFS", "INDEX-MATCH",
    "Analytical reporting", "Project documentation", "Git/GitHub", "Streamlit",
]


DOCS: list[DocSpec] = [
    DocSpec(
        "Master Resume",
        "Artificial Intelligence Graduate | Data Analysis | Applied AI | Project Coordination",
        "00_master",
        "Jinhyeok_Kim_Master_Resume.docx",
        "resume",
        "Artificial Intelligence graduate from UTS with practical experience across data analysis, applied machine learning, computer vision and project coordination. Skilled in turning messy datasets, unstructured text and video inputs into structured findings, working demos and clear reports.",
        MASTER_SKILLS,
        ["BioGeoDA", "AI Gym Trainer", "Tennis", "Project Management"],
        {k: v["bullets"] for k, v in PROJECTS.items()},
    ),
    DocSpec(
        "Master CV",
        "Artificial Intelligence Graduate | Applied AI Developer | Data Analyst",
        "00_master",
        "Jinhyeok_Kim_Master_CV.docx",
        "cv",
        "AI graduate with hands-on academic and portfolio experience in NLP, computer vision, data preparation, model evaluation, Excel-based planning and project documentation. The CV presents a broader record for graduate roles across data, AI, analytics and technology-enabled operations.",
        MASTER_SKILLS,
        ["BioGeoDA", "AI Gym Trainer", "Tennis", "Project Management"],
        {k: v["bullets"] for k, v in PROJECTS.items()},
        "Build a graduate career in applied AI and data analysis, using structured evidence, reliable reporting and practical software delivery to support business and technical teams.",
    ),
    DocSpec(
        "POSCO Associate Project Manager Resume",
        "Artificial Intelligence Graduate | Project Coordination | Data Analysis | Planning and Reporting",
        "01_posco_associate_project_manager",
        "Jinhyeok_Kim_POSCO_Associate_Project_Manager_Resume.docx",
        "resume",
        "Artificial Intelligence graduate with academic project leadership, data analysis, planning and reporting experience. Strong foundation in Excel-based scheduling, WBS development, risk identification, resource estimation and cross-functional communication.",
        ["Project coordination", "Task allocation", "WBS", "Gantt scheduling", "Critical path", "Risk management", "Budget planning", "Resource planning", "Microsoft Excel", "PowerPoint", "Analytical reporting", "Stakeholder communication", "Python", "Data analysis"],
        ["Project Management", "BioGeoDA", "AI Gym Trainer"],
        {
            "Project Management": [
                "Led a seven-member academic team through planning, task allocation, section review and final integration for a proposed AUD 1.5 million Event Management System.",
                "Coordinated WBS, AON and Excel-based Gantt planning with more than 100 tasks, dependencies, milestones, float and critical-path analysis.",
                "Reviewed risk, budget, human-resource, stakeholder and change-control sections to improve consistency before final submission.",
            ],
            "BioGeoDA": PROJECTS["BioGeoDA"]["bullets"][:3],
            "AI Gym Trainer": PROJECTS["AI Gym Trainer"]["bullets"][:2],
        },
    ),
    DocSpec(
        "POSCO Associate Project Manager CV",
        "Artificial Intelligence Graduate | Project Coordination | Planning and Reporting",
        "01_posco_associate_project_manager",
        "Jinhyeok_Kim_POSCO_Associate_Project_Manager_CV.docx",
        "cv",
        "AI graduate repositioned for project-support and associate project-management roles, combining data analysis, Excel planning, team leadership, documentation and structured problem solving.",
        ["Project coordination", "Academic team leadership", "Schedule tracking", "WBS", "AON", "Gantt charts", "Critical path", "Risk register", "Budget planning", "Human-resource estimation", "Stakeholder communication", "Microsoft Excel", "PowerPoint", "Python", "Data validation"],
        ["Project Management", "BioGeoDA", "AI Gym Trainer", "Tennis"],
        {},
        "Contribute to project planning, reporting and delivery support while applying data-analysis discipline to schedules, risks, resources and stakeholder information.",
    ),
    DocSpec(
        "Daesang Business Development Resume",
        "Artificial Intelligence Graduate | Business Analysis | Market Research | Data-Driven Decision Support",
        "02_daesang_business_development",
        "Jinhyeok_Kim_Daesang_Business_Development_Resume.docx",
        "resume",
        "AI graduate with data analysis, structured research, evidence review and Korean-English communication capability. Experienced in evaluating data sources, finding gaps, preparing reports and translating technical work into practical recommendations.",
        ["Business research", "Data-source evaluation", "Gap analysis", "Comparative analysis", "Structured reporting", "Data cleaning", "Microsoft Excel", "Charts", "PivotTables", "Python", "Pandas", "NLP", "Presentation preparation", "Korean-English communication", "Chinese HSK Level 4"],
        ["BioGeoDA", "Project Management", "AI Gym Trainer"],
        {
            "BioGeoDA": [
                "Analysed structured and unstructured plant-trait sources to identify information gaps and convert source evidence into usable trait records.",
                "Built NLP workflows using TF-IDF, Logistic Regression and BERT QA to compare keyword-based and context-based extraction methods.",
                "Contributed to team validation of 844 retained records from more than 2,400 generated candidates, supporting evidence-based reporting rather than unreviewed automation.",
            ],
            "Project Management": [
                "Led team coordination and report integration for an academic business-system planning scenario with budget, stakeholder, risk and resource components.",
                "Prepared planning outputs that translated broad business requirements into scope, milestones, risks and operational assumptions.",
            ],
            "AI Gym Trainer": PROJECTS["AI Gym Trainer"]["bullets"][:2],
        },
    ),
    DocSpec(
        "Daesang Business Development CV",
        "Artificial Intelligence Graduate | Business Analysis | Research and Reporting",
        "02_daesang_business_development",
        "Jinhyeok_Kim_Daesang_Business_Development_CV.docx",
        "cv",
        "AI graduate tailored for business development support, with strengths in data analysis, opportunity assessment, structured research, evidence-based communication and cross-cultural Korean-English collaboration.",
        ["Business analysis", "Market and source research", "Opportunity framing", "Gap analysis", "Comparative evaluation", "Data visualisation", "Microsoft Excel", "Charts", "PivotTables", "Python", "Pandas", "NLP", "Project documentation", "Korean native", "Chinese HSK Level 4"],
        ["BioGeoDA", "Project Management", "AI Gym Trainer", "Tennis"],
        {},
        "Support business-development decisions through disciplined research, clean analysis, clear reporting and strong Korean-English communication.",
    ),
    DocSpec(
        "Daesang SCM Coordinator Resume",
        "Artificial Intelligence Graduate | Data Analysis | Supply Chain Support | Process Coordination",
        "03_daesang_scm",
        "Jinhyeok_Kim_Daesang_SCM_Coordinator_Resume.docx",
        "resume",
        "AI and data graduate with transferable experience in data cleaning, validation, tracking, standardisation, Excel analysis and process-oriented reporting. Suitable for SCM support roles requiring accuracy, follow-up and structured information handling.",
        ["Microsoft Excel - Intermediate", "PivotTables", "XLOOKUP/VLOOKUP", "INDEX-MATCH", "SUMIFS/COUNTIFS", "Data cleaning", "Data validation", "Duplicate review", "Traceability", "Process monitoring", "Structured reporting", "Scheduling", "Python", "Pandas", "Korean-English communication"],
        ["BioGeoDA", "Project Management", "Tennis"],
        {
            "BioGeoDA": [
                "Standardised multi-source datasets by reviewing missing values, duplicate records, inconsistent species names and mixed trait formats.",
                "Maintained source, page, paragraph and context metadata so extracted records could be traced back to supporting evidence.",
                "Contributed to validation of 844 retained records from more than 2,400 generated candidates, reinforcing accuracy and review discipline.",
            ],
            "Project Management": [
                "Coordinated an Excel-based schedule and WBS with more than 100 tasks, dependencies, durations, milestones and resource estimates.",
                "Reviewed risk, vendor, stakeholder and budget-planning sections for consistency across a simulated 18-month delivery plan.",
            ],
            "Tennis": PROJECTS["Tennis"]["bullets"][:2],
        },
    ),
    DocSpec(
        "Daesang SCM Coordinator CV",
        "Artificial Intelligence Graduate | Data Analysis | SCM Support | Process Coordination",
        "03_daesang_scm",
        "Jinhyeok_Kim_Daesang_SCM_Coordinator_CV.docx",
        "cv",
        "AI and data graduate tailored for SCM coordination, combining dataset maintenance, validation, Excel reporting, process monitoring and analytical problem solving. Experience is academic and portfolio-based, with clear transferability to supply-chain support workflows.",
        ["Data maintenance", "Data validation", "Standardisation", "Traceability", "Microsoft Excel", "PivotTables", "Lookup functions", "SUMIFS/COUNTIFS", "Scheduling", "Process monitoring", "Exception identification", "Reporting", "Python", "Pandas", "Project coordination"],
        ["BioGeoDA", "Project Management", "Tennis", "AI Gym Trainer"],
        {},
        "Apply data-quality discipline and process coordination skills to supply-chain support, reporting and operational follow-up.",
    ),
    DocSpec(
        "Kumho Tyre Logistics Coordinator Resume",
        "Artificial Intelligence Graduate | Data Reporting | Logistics Support | Process Monitoring",
        "04_kumho_logistics",
        "Jinhyeok_Kim_Kumho_Tyre_Logistics_Coordinator_Resume.docx",
        "resume",
        "Data-oriented AI graduate with Excel, reporting, process monitoring, data maintenance and project coordination skills. Brings careful evidence handling, Korean-English communication and structured problem solving for logistics support environments.",
        ["Excel reporting", "Data reconciliation concepts", "Data maintenance", "Process monitoring", "Exception identification", "Cost and resource estimation", "Schedule monitoring", "Vendor coordination planning", "Microsoft Excel", "PivotTables", "Lookup functions", "Python", "Pandas", "Korean-English communication"],
        ["BioGeoDA", "Project Management", "Tennis"],
        {
            "BioGeoDA": [
                "Maintained traceable data outputs by preserving source, page, paragraph and context metadata through extraction and validation workflows.",
                "Cleaned and standardised multi-source records by checking missing values, duplicates, naming inconsistency and mixed trait formats.",
                "Used structured review to support reliable reporting from more than 2,400 generated candidate records.",
            ],
            "Project Management": [
                "Coordinated schedule, cost and resource estimates across more than 100 tasks in an academic logistics-heavy event-system planning scenario.",
                "Reviewed vendor coordination, site preparation, risk, budget and closeout components within an 18-month simulated delivery plan.",
            ],
            "Tennis": PROJECTS["Tennis"]["bullets"][:2],
        },
    ),
    DocSpec(
        "Kumho Tyre Logistics Coordinator CV",
        "Artificial Intelligence Graduate | Data Reporting | Logistics Support | Process Monitoring",
        "04_kumho_logistics",
        "Jinhyeok_Kim_Kumho_Tyre_Logistics_Coordinator_CV.docx",
        "cv",
        "Data-oriented AI graduate tailored for logistics coordination, with transferable strengths in Excel reporting, data maintenance, process monitoring, schedule awareness and careful exception identification.",
        ["Logistics support", "Data reporting", "Data maintenance", "Process monitoring", "Schedule tracking", "Cost and resource estimation", "Vendor coordination planning", "Microsoft Excel", "PivotTables", "Lookup functions", "Python", "Pandas", "Korean native", "English professional communication"],
        ["BioGeoDA", "Project Management", "Tennis", "AI Gym Trainer"],
        {},
        "Support logistics operations through accurate data maintenance, clear reporting, disciplined follow-up and fast learning of enterprise platforms.",
    ),
]


def set_run_font(run, size=10, bold=False, color="1f2937"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=4, line=1.05):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bottom_border(paragraph, color="CBD5E1"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def prepare_doc(kind: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55 if kind == "resume" else 0.65)
    section.bottom_margin = Inches(0.55 if kind == "resume" else 0.65)
    section.left_margin = Inches(0.62 if kind == "resume" else 0.72)
    section.right_margin = Inches(0.62 if kind == "resume" else 0.72)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9.6 if kind == "resume" else 10.2)
    for style_name, size, colour in [
        ("Heading 1", 13 if kind == "resume" else 14, "0F766E"),
        ("Heading 2", 11.2 if kind == "resume" else 12, "334155"),
        ("Heading 3", 10.5 if kind == "resume" else 11, "334155"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(7 if kind == "resume" else 9)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True
    return doc


def add_header(doc: Document, spec: DocSpec):
    p = doc.add_paragraph()
    set_spacing(p, after=1)
    run = p.add_run(CONTACT["name"])
    set_run_font(run, 21, True, "0F172A")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph()
    set_spacing(sub, after=2)
    r = sub.add_run(spec.subtitle)
    set_run_font(r, 10.2, True, "0F766E")
    contact = doc.add_paragraph()
    set_spacing(contact, after=6)
    r = contact.add_run(
        f'{CONTACT["location"]} | {CONTACT["phone"]} | {CONTACT["email"]} | Portfolio: {CONTACT["portfolio"]} | GitHub: {CONTACT["github"]}'
    )
    set_run_font(r, 8.7, False, "475569")
    add_bottom_border(contact)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_para(doc, text, size=None, bold=False, colour="1f2937", after=4):
    p = doc.add_paragraph()
    set_spacing(p, after=after)
    r = p.add_run(text)
    set_run_font(r, size or 9.6, bold, colour)
    return p


def add_bullet(doc, text, size=None):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=2, line=1.02)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    r = p.add_run(text)
    set_run_font(r, size or 9.35)
    return p


def add_skill_block(doc, skills, kind):
    max_per_line = 6 if kind == "resume" else 5
    for i in range(0, len(skills), max_per_line):
        add_para(doc, " | ".join(skills[i : i + max_per_line]), 9.2 if kind == "resume" else 9.7, after=2)


def add_education(doc, kind):
    for item in EDUCATION:
        p = doc.add_paragraph()
        set_spacing(p, after=1)
        r = p.add_run(f'{item["degree"]} - {item["period"]}')
        set_run_font(r, 10 if kind == "resume" else 10.5, True, "0F172A")
        add_para(doc, item["school"], 9.2 if kind == "resume" else 9.8, colour="475569", after=1)
        for detail in item["details"]:
            add_bullet(doc, detail, 9.1 if kind == "resume" else 9.6)


def add_project(doc, key, bullets, kind, cv_detail=False):
    project = PROJECTS[key]
    p = doc.add_paragraph()
    set_spacing(p, before=3, after=1)
    r = p.add_run(f'{project["title"]} | {project["period"]}')
    set_run_font(r, 10 if kind == "resume" else 10.7, True, "0F172A")
    add_para(doc, f'{project["context"]} | {project["tools"]}', 8.8 if kind == "resume" else 9.3, True, "475569", after=2)
    if cv_detail:
        intro = {
            "BioGeoDA": "Problem: botanical traits were scattered across structured datasets and OCR-ready journal text. Approach: combine cleaning, standardisation, NLP baselines and validation to make outputs reviewable.",
            "AI Gym Trainer": "Problem: exercise videos need interpretable analysis rather than a single opaque classification label. Approach: combine video preprocessing, model inference, pose landmarks and structured reporting.",
            "Tennis": "Problem: a tennis ball is small, fast and frequently obscured in broadcast footage. Approach: use classical detection signals plus Kalman smoothing and evaluation diagnostics.",
            "Project Management": "Problem: a complex event-system scenario needed a credible plan rather than a built product. Approach: lead planning, schedule, scope, resource, risk and budget documentation for an academic simulation.",
        }[key]
        add_para(doc, intro, 9.6, after=3)
    for bullet in bullets:
        add_bullet(doc, bullet, 9.15 if kind == "resume" else 9.55)


def write_doc(spec: DocSpec):
    doc = prepare_doc(spec.kind)
    add_header(doc, spec)
    add_heading(doc, "Professional Summary" if spec.kind == "resume" else "Career Profile")
    add_para(doc, spec.summary, 9.5 if spec.kind == "resume" else 10.0, after=4)
    if spec.cv_objective:
        add_heading(doc, "Career Objective", 2)
        add_para(doc, spec.cv_objective, 9.8, after=4)
    add_heading(doc, "Core Skills" if spec.kind == "resume" else "Technical and Business Competencies")
    add_skill_block(doc, spec.skills, spec.kind)
    add_heading(doc, "Education")
    add_education(doc, spec.kind)
    if spec.kind == "cv":
        add_heading(doc, "Relevant Coursework")
        add_para(doc, "; ".join(COURSEWORK), 9.5, after=3)
        add_heading(doc, "Applied Academic Evidence", 2)
        academic_evidence = [
            "Data analytics coursework covered cleaning, validation, statistical interpretation and communication of findings from structured datasets.",
            "Advanced analytics and machine-learning study developed model-comparison discipline, feature preparation, evaluation and limitation reporting.",
            "Design, Data, and Decisions strengthened Excel-based analysis using PivotTables, lookup functions, conditional formatting, charts and practical decision support.",
            "Project Management and the Professional developed WBS, AON, Gantt scheduling, risk, stakeholder, budget and human-resource planning capability.",
            "Communication for IT Professionals supported clear technical documentation, report structure and audience-aware presentation of analysis.",
        ]
        for item in academic_evidence:
            add_bullet(doc, item, 9.55)
    add_heading(doc, "Selected Project Experience" if spec.kind == "resume" else "Detailed Project Experience")
    for key in spec.project_order:
        bullets = spec.project_bullets.get(key) or PROJECTS[key]["bullets"]
        max_bullets = 3 if spec.kind == "resume" else 4
        add_project(doc, key, bullets[:max_bullets], spec.kind, spec.kind == "cv")
    if spec.kind == "cv":
        add_heading(doc, "Leadership and Collaboration")
        for item in [
            "Led a seven-member academic team by setting direction, assigning tasks, reviewing sections and integrating the final project-management report.",
            "Worked in a four-person UTS capstone team while separately publishing a portfolio-safe reconstruction of the AI/NLP contribution.",
            "Communicated technical project results through reports, README documentation, demos and portfolio case-study writing.",
        ]:
            add_bullet(doc, item, 9.55)
        add_heading(doc, "Role Fit and Transferable Strengths")
        role_fit = [
            f"Can contribute through {spec.skills[0].lower()}, {spec.skills[1].lower()} and {spec.skills[2].lower()} while continuing to build professional workplace experience.",
            "Comfortable working from imperfect information, checking source evidence and separating verified facts from assumptions.",
            "Able to translate technical outputs into concise documentation for business, operational and non-technical stakeholders.",
            "Brings Korean native language capability, professional English communication and HSK Level 4 Chinese for cross-cultural team environments.",
        ]
        if "SCM" in spec.filename or "Logistics" in spec.filename:
            role_fit.append(
                "Although direct SAP experience was not verified, the project record shows fast learning across structured data systems, web applications and analytical workflows."
            )
        else:
            role_fit.append(
                "Open to learning role-specific platforms quickly, with a project record that shows practical adaptation across datasets, codebases and reporting requirements."
            )
        for item in role_fit:
            add_bullet(doc, item, 9.55)
    add_heading(doc, "Languages")
    for item in LANGUAGES:
        add_bullet(doc, item, 9.15 if spec.kind == "resume" else 9.55)
    add_heading(doc, "Work Rights and Availability")
    add_para(doc, WORK_RIGHTS, 9.1 if spec.kind == "resume" else 9.5, after=2)
    out_dir = CAREER / spec.folder
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / spec.filename
    doc.save(out_path)
    return out_path


def convert_to_pdf(docx_path: Path):
    if not SOFFICE.exists():
        return None
    with tempfile.TemporaryDirectory(dir="/private/tmp") as tmp:
        profile = Path(tmp) / "lo-profile"
        profile.mkdir(parents=True, exist_ok=True)
        env = {
            **os.environ,
            "HOME": str(profile),
            "TMPDIR": "/private/tmp",
            "TEMP": "/private/tmp",
            "TMP": "/private/tmp",
            "XDG_CONFIG_HOME": str(profile / "xdg_config"),
            "XDG_CACHE_HOME": str(profile / "xdg_cache"),
        }
        Path(env["XDG_CONFIG_HOME"]).mkdir(parents=True, exist_ok=True)
        Path(env["XDG_CACHE_HOME"]).mkdir(parents=True, exist_ok=True)
        subprocess.run(
            [
                str(SOFFICE),
                f"-env:UserInstallation=file://{profile}",
                "--invisible",
                "--headless",
                "--norestore",
                "--convert-to",
                "pdf",
                "--outdir",
                str(docx_path.parent),
                str(docx_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            env=env,
        )
    return docx_path.with_suffix(".pdf")


def write_text(path: Path, content: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content.strip() + "\n", encoding="utf-8")


def write_source_docs():
    source = CAREER / "source"
    write_text(
        source / "verified_candidate_facts.md",
        """
# Verified Candidate Facts

| Fact | Source | Status | Notes |
|---|---|---|---|
| Full name: Jinhyeok Kim | Existing Pages resume preview | Verified | Header reads JINHYEOK KIM. |
| Location: Sydney, NSW, Australia | Existing Pages resume preview | Verified | Header line. |
| Mobile: 0468 544 027 | Existing Pages resume preview | Verified | Header line. |
| Email: winsuwon127@gmail.com | Existing Pages resume preview and portfolio source | Verified | Header line and site contact data. |
| Portfolio URL: jinhyeok-portfolio-amber.vercel.app | Existing Pages resume preview and README | Verified | Header line and README. |
| GitHub: github.com/Jinhyeok513 | Existing Pages resume preview and project repositories | Verified | Header line and project URLs. |
| Bachelor of Artificial Intelligence, UTS | Existing Pages resume preview | Verified | Education section. |
| Bachelor completed 2026 | Existing Pages resume preview | Verified | Education section. |
| Graduated with Distinction; WAM 77.94; GPA 5.88/7.00 | Existing Pages resume preview | Verified | Education bullet. |
| Diploma of Information Technology, UTS College | Existing Pages resume preview | Verified | Education section. |
| Diploma year: 2023 | Existing Pages resume preview | Verified | Education section. |
| Relevant coursework list | Desktop/ASSINGMENT/Courses.pdf | Verified | Canvas course export dated 03/08/2026. |
| Korean native | Existing Pages resume preview text extraction | Verified | Professional summary text. |
| English professional communication skills | Existing Pages resume preview text extraction | Verified | Professional summary text. |
| Chinese HSK Level 4 | User-provided prompt | User-stated | No local certificate found. |
| Current visa subclass, exact work-rights wording and full-time availability date | Existing resume preview/text extraction | Not verified | The Pages file exposes a Work Right heading in binary strings, but the actual wording was not readable. Confirm before final submission. |
""",
    )
    write_text(
        source / "project_evidence.md",
        """
# Project Evidence

| Claim | Source file | Source location | Verified | Notes |
|---|---|---|---|---|
| BioGeoDA uses a 1,674-species checklist | /Users/jin/Desktop/ASSINGMENT/BioGioDa/acp_a3_Project4GroupA.docx | ALA checklist discussion and Figure 8 caption | Yes | Team capstone result. |
| BioGeoDA generated more than 2,400 AI candidates and retained 844 records after manual comparison | /Users/jin/Desktop/ASSINGMENT/BioGioDa/acp_a3_Project4GroupA.docx | Section 5.2 Key Insights from AI | Yes | Team validation result; not solo ownership. |
| BioGeoDA TF-IDF Logistic Regression accuracy 0.9071 and macro F1 0.4617 | /Users/jin/Desktop/ASSINGMENT/BioGioDa/acp_a3_Project4GroupA.docx and BioGeoDA README | AI model performance section | Yes | Used with class-imbalance limitation. |
| BioGeoDA public repository is portfolio-safe reconstruction | /Users/jin/Desktop/File/jinhyeok-portfolio/projects/code/BioGeoDA/README.md | Original Capstone Context and Data and Copyright | Yes | Excludes private source data and teammate-owned code. |
| AI Gym Trainer uses MobileNetV2, MediaPipe Pose, OpenCV and CSV/JSON exports | /Users/jin/Desktop/File/jinhyeok-portfolio/projects/code/gym-ai-trainer/README.md | Features, AI/ML Pipeline and Tech Stack | Yes | Portfolio app. |
| AI Gym Trainer MobileNetV2 reached 72.0% video accuracy and 78.18% macro F1 | /Users/jin/Desktop/File/jinhyeok-portfolio/projects/code/gym-ai-trainer/MODEL_CARD.md | Training And Evaluation Evidence | Yes | Include limitations. |
| AI Gym Trainer 10-video field test reached 20% video-level accuracy | /Users/jin/Desktop/File/jinhyeok-portfolio/projects/code/gym-ai-trainer/MODEL_CARD.md | Real-World Validation | Yes | Used as domain-shift evidence. |
| Tennis project uses Kalman filtering, OpenCV, evaluation metrics and 207-frame output | /Users/jin/Desktop/jinhyeok-portfolio/src/data/portfolio.ts | Tennis project entry | Verified from existing portfolio data | No local Tennis source repo was found in this workspace; keep wording conservative. |
| Tennis sparse-label metrics: mean IoU 0.1827, Success@IoU>=0.5 2.4%, mean CLE 131.18 px | /Users/jin/Desktop/jinhyeok-portfolio/src/data/portfolio.ts | Tennis project keyResults | Verified from existing portfolio data | Preserve as reported metrics. |
| Event Management System report was a seven-member academic team assignment | /Users/jin/Desktop/ASSINGMENT/Project Manage/Project report.pdf | Coversheet | Yes | Do not publish personal details from original. |
| EMS scenario proposed AUD 1.5M budget and 30 Nov 2026 completion target | /Users/jin/Desktop/ASSINGMENT/Project Manage/Project report.pdf | Executive summary and success criteria | Yes | Academic simulation. |
| EMS WBS/Gantt has more than 100 tasks and an 18-month-style schedule | /Users/jin/Desktop/ASSINGMENT/Project Manage/WBS_List .csv and Gantt Chart .xlsx | CSV task rows and Sheet1 | Yes | 150 rows including summaries; task IDs exceed 100. |
| Stock assignment exists but is not portfolio-ready | /Users/jin/Desktop/ASSINGMENT/Stock | Notebook, report and outputs | Yes | Used only as internal evidence of ML coursework; removed from featured portfolio. |
""",
    )
    write_text(
        source / "role_keyword_mapping.md",
        """
# Role Keyword Mapping

| Target role | Primary positioning | Keywords emphasised | Claims avoided |
|---|---|---|---|
| Master | AI graduate for data, applied AI and project support | Python, data cleaning, NLP, computer vision, Excel, reporting, GitHub, project documentation | Direct professional experience not supported by files |
| POSCO Associate Project Manager | AI graduate with academic project leadership and planning exposure | WBS, AON, Gantt, critical path, risk, budget, resources, milestone tracking, stakeholder communication | Professional project manager employment |
| Daesang New Business Development | AI graduate with research, analysis and Korean-English communication | Business analysis, source evaluation, gap analysis, comparative analysis, structured reporting, presentation preparation | Sales revenue, client acquisition, market-entry wins |
| Daesang SCM Coordinator | Data graduate with process and data-quality transferability | Excel, lookup functions, PivotTables, data cleaning, duplicate review, traceability, scheduling, process monitoring | Inventory management, demand planning or SAP experience |
| Kumho Tyre Logistics Coordinator | Data reporting and logistics-support transferability | Data maintenance, process monitoring, schedule tracking, cost/resource estimation, vendor planning, exception identification | SAP experience or direct logistics employment |
""",
    )
    write_text(
        source / "document_change_log.md",
        """
# Document Change Log

## Files created

- 10 editable DOCX files and 10 selectable-text PDFs under docs/career.
- 4 portfolio case-study Markdown files under docs/career/portfolio_case_studies.
- 4 source/audit Markdown files under docs/career/source.
- Public resume PDF copied to public/resume/Jinhyeok_Kim_Master_Resume.pdf.

## Claims added

- BioGeoDA 1,674-species checklist, 2,400+ AI candidates and 844 retained records.
- BioGeoDA TF-IDF Logistic Regression accuracy 0.9071 and macro F1 0.4617.
- AI Gym Trainer MobileNetV2 72.0% video accuracy, 78.18% macro F1 and 20% field-test accuracy.
- EMS academic case-study scale: seven-member team, AUD 1.5M proposed budget, 100+ tasks and 18-month plan.

## Claims removed or constrained

- Removed Stock Market Prediction from featured portfolio and Resume/CV selected projects.
- Removed LinkedIn from site contact data and employment documents.
- Avoided Advanced Excel, Solver, Goal Seek, SAP and direct supply-chain/logistics employment claims.
- Reworded team-level results as contributed/team outcomes rather than solo ownership.

## Facts requiring later confirmation

- Current visa subclass, exact Australian work-rights wording and full-time availability date were not readable from the Pages file preview or extracted strings. Confirm before sending applications.

## Portfolio changes

- Hero and contact copy now presents Jinhyeok as an AI/data graduate with practical applied AI projects.
- Stock project card removed.
- Event Management System academic case study added as the fourth project.
- Resume download button now points to the generated master resume PDF.
- Tennis video remains embedded as an internal playable portfolio video.
""",
    )


def write_case_studies():
    target = CAREER / "portfolio_case_studies"
    public_target = ROOT / "public" / "case-studies"
    write_text(
        target / "biogeoda.md",
        """
# BioGeoDA - Australian Plant Data Integration and AI-Assisted Trait Extraction

## Overview
BioGeoDA is a UTS capstone-origin AI/data project focused on extracting plant trait information from structured datasets and OCR-ready botanical text. The public repository is a portfolio-safe reconstruction of Jinhyeok Kim's AI/NLP contribution.

## Problem
Plant traits often appear in inconsistent formats across formal databases and journal-style descriptions. The work needed cleaning, standardisation, extraction and validation so that trait information could be reviewed rather than treated as a black-box model output.

## My Role
Jinhyeok contributed across data preprocessing, trait-to-value mappings, QA example generation, TF-IDF baseline work, BERT QA integration, propagation post-processing, validation support and public demo restructuring.

## Scope
The project worked with a 1,674-species checklist. Team validation reviewed more than 2,400 AI-generated candidates and retained 844 correctly matched trait records.

## Approach
- Cleaned missing values, duplicate records, species naming inconsistencies and mixed trait formats.
- Used keyword tagging, fuzzy matching and metadata preservation for source traceability.
- Compared TF-IDF Logistic Regression with BERT-based extractive question answering.
- Rebuilt the public demo with synthetic sample data to avoid exposing private APJ material.

## Results
- TF-IDF Logistic Regression recorded 90.7% accuracy and 46.2% macro F1 in historical capstone notebooks.
- BERT QA recorded an evaluation loss of approximately 0.321 in the project documentation.
- Team validation retained 844 records after manual comparison.

## Limitations
The public demo is not a benchmark. It excludes private data, original dashboards, full OCR output and teammate-owned code. Confidence calibration was not implemented.

## Tools
Python, Pandas, Streamlit, scikit-learn, TF-IDF, Logistic Regression, BERT QA, GitHub.
""",
    )
    write_text(
        target / "ai_gym_trainer.md",
        """
# AI Gym Trainer - Exercise Video Analysis Application

## Overview
AI Gym Trainer is an end-to-end portfolio application for short exercise-video analysis. It combines exercise classification, pose landmark extraction, movement metrics, annotated output video and downloadable reporting.

## Problem
Workout video analysis needs more than a single prediction label. A useful review flow should show what the system saw, where it was uncertain and how results were generated.

## My Role
Jinhyeok built and documented the application workflow from video upload through preprocessing, MobileNetV2 inference, MediaPipe Pose analysis, OpenCV rendering and CSV/JSON reporting.

## Approach
- Accepted a short video clip through a web interface.
- Sampled frames and classified supported exercise types with MobileNetV2.
- Extracted body landmarks with MediaPipe Pose.
- Calculated joint-angle metrics and rule-based feedback.
- Returned annotated video, session metrics and downloadable frame-level data.

## Results
- MobileNetV2 recorded 72.0% video accuracy and 78.18% video macro F1 on the documented held-out comparison data.
- A 10-video field test reached 20% video-level accuracy, documenting domain shift honestly.
- The pipeline supports annotated browser-playable video plus CSV and JSON exports.

## Limitations
This is a research prototype, not a medical or professional coaching tool. Confidence is not fully calibrated, rep counts remain heuristic and diverse user-recorded videos would be needed for robust public use.

## Tools
Next.js, React, TypeScript, FastAPI, Python, TensorFlow, MobileNetV2, MediaPipe Pose, OpenCV, Pandas, NumPy.
""",
    )
    write_text(
        target / "tennis_ball_tracking.md",
        """
# Tennis Ball Tracking and Trajectory Visualisation

## Overview
This project demonstrates classical computer-vision tracking of a tennis ball in broadcast-style footage. The portfolio presents the result as a video-based tracking and trajectory visualisation project rather than overstating fully reliable real-time performance.

## Problem
The tennis ball is small, fast, blurred and frequently hidden by court lines, players, net regions and camera motion. A tracking pipeline must handle noise and short detection gaps.

## My Role
Jinhyeok implemented and documented the Kalman-filter tracking workflow, detection cues, trajectory rendering and evaluation diagnostics presented in the portfolio.

## Approach
- Used OpenCV frame processing with HSV colour cues, whiteness cues, optical flow and blobness scoring.
- Applied court masking and net-region suppression to reduce false positives.
- Used Kalman filtering to smooth noisy candidate detections and predict short gaps.
- Packaged a portfolio-ready MP4 result video and frame-level diagnostics.

## Results
- Published a 207-frame tracking output with representative captures and trajectory visualisation.
- Reported sparse-label metrics: mean IoU 0.1827, Success@IoU>=0.5 of 2.4% and mean CLE 131.18 px.

## Limitations
The project is best presented as video-based tracking and trajectory visualisation. The sparse metrics show that the method is educational and diagnostic, not a production sports analytics system.

## Tools
Python, OpenCV, NumPy, Kalman filtering, FFmpeg, IoU/CLE evaluation.
""",
    )
    write_text(
        target / "project_management_case_study.md",
        """
# Event Management System Project Plan - Team Lead

## Overview
This is a sanitised academic Project Management case study based on a UTS team assignment. It describes a proposed Event Management System for a major Sydney festival scenario. The original 68-page report and Excel workbook are not public because they contain student details and jointly authored assessment material.

## Problem
The assignment scenario required a credible project plan for a complex event technology system. The team needed to define scope, stakeholders, risks, budget, schedule, resourcing and quality controls without claiming the system was actually commissioned or built.

## My Role
Jinhyeok led the seven-member academic team. He set direction, assigned tasks, guided what each member should produce, reviewed sections, corrected inconsistencies and integrated the final report for submission.

## Scope
- Proposed AUD 1.5 million budget.
- 18-month delivery plan ending 30 November 2026.
- More than 100 tasks across WBS and Gantt planning.
- WBS, AON dependency analysis, Gantt schedule and critical-path review.
- Stakeholder, communication, risk, scope, quality, budget and resource planning.

## Approach
- Broke the project into five major phases: initiation/planning, design/development, testing/validation, deployment/execution and closeout/reporting.
- Built a task list with dependencies, durations, milestones, float and critical-path indicators.
- Reviewed budget and human-resource assumptions against the simulated delivery constraints.
- Sanitised public portfolio wording to remove private student information and assessment-specific pages.

## Results
The public case study demonstrates project coordination, Excel-based planning, structured reporting, risk thinking and leadership in an academic team setting. It does not claim professional project-manager employment or real client delivery.

## Limitations
This was an academic simulation. The EMS was planned, not built or commissioned by a real client.

## Tools
Microsoft Excel, WBS, AON, Gantt scheduling, risk register, budget planning, resource estimation, PowerPoint, report integration.
""",
    )
    if public_target.exists():
        shutil.rmtree(public_target)
    public_target.mkdir(parents=True, exist_ok=True)
    for path in target.glob("*.md"):
        shutil.copy2(path, public_target / path.name)


def make_pm_image():
    out = ROOT / "public" / "images" / "projects" / "project-management-case-study.png"
    out.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1400, 900), "#08111f")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 54)
        head_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 34)
        body_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 28)
        small_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 23)
    except Exception:
        title_font = head_font = body_font = small_font = ImageFont.load_default()
    draw.rectangle((0, 0, 1400, 900), fill="#08111f")
    draw.rectangle((70, 70, 1330, 830), outline="#1dd3b0", width=4)
    draw.text((100, 105), "Academic Project Management Case Study", fill="#e6fffb", font=title_font)
    draw.text((100, 180), "Event Management System Project Plan - Team Lead", fill="#a7f3d0", font=body_font)
    cards = [
        ("AUD 1.5M", "proposed budget"),
        ("18 months", "delivery plan"),
        ("100+ tasks", "WBS + Gantt"),
        ("7 members", "academic team"),
        ("5 phases", "planning to closeout"),
        ("5% contingency", "budget buffer"),
    ]
    x0, y0 = 100, 270
    for i, (value, label) in enumerate(cards):
        x = x0 + (i % 3) * 400
        y = y0 + (i // 3) * 160
        draw.rounded_rectangle((x, y, x + 340, y + 115), radius=18, fill="#10243a", outline="#2dd4bf", width=2)
        draw.text((x + 28, y + 24), value, fill="#ffffff", font=head_font)
        draw.text((x + 28, y + 72), label, fill="#b6c8d8", font=small_font)
    phases = ["Initiation", "Design", "Testing", "Deployment", "Closeout"]
    y = 660
    for i, phase in enumerate(phases):
        x = 112 + i * 240
        draw.rounded_rectangle((x, y, x + 180, y + 58), radius=14, fill="#0f766e")
        draw.text((x + 22, y + 16), phase, fill="#ffffff", font=small_font)
        if i < len(phases) - 1:
            draw.line((x + 182, y + 29, x + 230, y + 29), fill="#7dd3fc", width=4)
            draw.polygon([(x + 230, y + 29), (x + 216, y + 20), (x + 216, y + 38)], fill="#7dd3fc")
    draw.text((100, 765), "Sanitised public summary - original report and workbook not published", fill="#94a3b8", font=small_font)
    img.save(out)
    return out


def main():
    if CAREER.exists():
        shutil.rmtree(CAREER)
    CAREER.mkdir(parents=True, exist_ok=True)
    PUBLIC_RESUME.mkdir(parents=True, exist_ok=True)
    write_source_docs()
    write_case_studies()
    make_pm_image()
    generated = []
    for spec in DOCS:
        docx_path = write_doc(spec)
        generated.append(docx_path)
        pdf_path = convert_to_pdf(docx_path)
        if pdf_path:
            generated.append(pdf_path)
    master_pdf = CAREER / "00_master" / "Jinhyeok_Kim_Master_Resume.pdf"
    if master_pdf.exists():
        shutil.copy2(master_pdf, PUBLIC_RESUME / "Jinhyeok_Kim_Master_Resume.pdf")
    print("\n".join(str(p.relative_to(ROOT)) for p in generated))


if __name__ == "__main__":
    main()
